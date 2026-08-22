# -*- coding: utf-8 -*-
"""Regenerate submission docs (title page, cover letter, data availability, highlights)."""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

AUTHORS = ("Zhipeng Wang", "Luning Wang", "Changsong Wang", "Pengli Zhai",
           "Zejun Liu", "Hui Feng", "Hongmei Liu", "Qian Hou", "Ming Guo")
AFFILS = [
    "1 Department of TCM, Zhongda Hospital, Southeast University, China",
    "2 Department of Rehabilitation Medicine, Zhongda Hospital, Southeast University, China",
    "3 Jiangbei Campus, Jiangsu Provincial Traditional Chinese Medicine Hospital, China",
]
CORR = "Changsong Wang, Department of TCM, Zhongda Hospital, Southeast University, China. Email: 101005664@seu.edu.cn"
TITLE = ("Path-AGNN-Cox: a reproducible statistical framework for testing "
         "patient-specific pathway rewiring in cancer survival analysis")
KEYWORDS = "survival analysis; graph neural network; pathway constraint; adaptive graph learning; cancer prognosis"

def base():
    d = Document()
    st = d.styles["Normal"]
    st.font.name = "Times New Roman"
    st.font.size = Pt(11)
    return d

def add(d, text, bold=False, align=None):
    p = d.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    if align: p.alignment = align
    return p

# ---------------- title page ----------------
d = base()
add(d, "Title Page", bold=True)
add(d, TITLE)
add(d, "")
add(d, "Authors", bold=True)
add(d, "Zhipeng Wang[1,*], Luning Wang[2,*], Changsong Wang[1,†], Pengli Zhai[3], Zejun Liu[1], Hui Feng[1], Hongmei Liu[1], Qian Hou[1], Ming Guo[1]")
add(d, "")
for a in AFFILS:
    add(d, a)
add(d, "* These authors contributed equally to this work and are co-first authors.")
add(d, "† Corresponding author: " + CORR)
add(d, "")
add(d, "Keywords", bold=True)
add(d, KEYWORDS)
add(d, "")
add(d, "Manuscript statistics", bold=True)
add(d, "- Abstract: 227 words (Background / Methods / Results / Conclusion)")
add(d, "- Main text: approximately 7,000 words (Introduction to Conclusion)")
add(d, "- Figures: 8 (Figures 1-8)")
add(d, "- Tables: 9 (Tables 1-9)")
add(d, "- Supplementary: none (all methods and tables integrated in the main text)")
add(d, "")
add(d, "Ethics approval", bold=True)
add(d, "This study was approved by the Ethics Committee of Zhongda Hospital Affiliated with Southeast University. All analyses used publicly available de-identified data; no new clinical data or patient samples were collected.")
add(d, "")
add(d, "Conflict of interest statement", bold=True)
add(d, "The authors declare no competing interests.")
add(d, "")
add(d, "Code and data availability", bold=True)
add(d, "All code is publicly available at https://github.com/wangzhipeng-1/Path-AGNN-Cox (MIT license) and installable via PyPI as path-agnn-cox (pip install path-agnn-cox). An archived snapshot is deposited on Zenodo. All data are from public repositories (TCGA via UCSC Xena and GDC; GEO; IMvigor210); per-patient data are not redistributed to protect privacy, and aggregate results required to reproduce every table and figure are provided in the repository.")
d.save(r"manuscript/title_page.docx")
print("title_page.docx saved")

# ---------------- cover letter ----------------
d = base()
add(d, "Dear Editors,", )
add(d, "")
add(d, "We are pleased to submit our manuscript, \u201c" + TITLE + "\u201d, for consideration.")
add(d, "")
add(d, "What the manuscript contributes.", bold=True)
add(d, "Transcriptomic survival models based on graph neural networks (GNNs) typically treat the underlying gene interaction topology as fixed and identical across patients. Path-AGNN-Cox constrains message passing to KEGG pathway subgraphs and learns per-patient pathway graphs; we additionally provide a statistical framework for formally testing patient-specific pathway rewiring, with edge- and pathway-level tests under false-discovery-rate control, label-permutation nulls, a static-model negative control, and a standard-GAT architectural control.")
add(d, "")
add(d, "Key points.", bold=True)
add(d, "1. A model that produces patient-specific pathway graphs, coupled to a Cox survival objective with dual regularization.")
add(d, "2. A statistical testing framework for pathway rewiring, benchmarked across 11 TCGA cancer types and 25 independent GEO cohorts, with three ablations and multiple controls.")
add(d, "3. Clinical and external anchoring: rewiring magnitude correlates with Ki-67 in two independent cohorts and with tumor mutational burden in LUAD, and is absent by construction in static pathway models.")
add(d, "")
add(d, "Reproducibility.", bold=True)
add(d, "All code is released as an open-source Python package with a pip-installable distribution, example notebooks, and reproducible pipelines covering every table and figure. No patient-level data are redistributed; aggregate outputs required for reproduction are provided.")
add(d, "")
add(d, "All authors have read and approved the manuscript, and none of the authors has a competing interest. This study was approved by the Ethics Committee of Zhongda Hospital Affiliated with Southeast University. The manuscript is not under consideration elsewhere.")
add(d, "")
add(d, "Thank you for your consideration.")
add(d, "")
add(d, "Sincerely,")
add(d, "Zhipeng Wang[1,*], Luning Wang[2,*], Changsong Wang[1,†], Pengli Zhai[3], Zejun Liu[1], Hui Feng[1], Hongmei Liu[1], Qian Hou[1], Ming Guo[1]")
for a in AFFILS:
    add(d, a)
add(d, "* These authors contributed equally to this work and are co-first authors.")
add(d, "† Corresponding author: " + CORR)
d.save(r"manuscript/cover_letter.docx")
print("cover_letter.docx saved")

# ---------------- data availability ----------------
d = base()
add(d, "Data Availability Statement", bold=True)
add(d, "All data used in this study are publicly available from their original sources. TCGA RNA-seq and clinical annotations were obtained from UCSC Xena and GDC; GEO series matrices and platform annotations were obtained from NCBI GEO; IMvigor210 data are available from the original publication (Mariathasan et al., Nature, 2018). Accession numbers are listed in Table 1 of the manuscript.")
add(d, "")
add(d, "To protect patient privacy, per-patient expression matrices, clinical records, and per-patient model outputs are not redistributed. Aggregate results required to reproduce every table and figure are provided in the repository, together with the complete analysis pipelines.")
add(d, "")
add(d, "No new biological data were generated in this study.")
d.save(r"manuscript/data_availability.docx")
print("data_availability.docx saved")

# ---------------- highlights ----------------
d = base()
add(d, "Highlights", bold=True)
add(d, "- Path-AGNN-Cox learns per-patient pathway graphs within KEGG pathway modules and links them to survival through a Cox objective.")
add(d, "- A reproducible statistical framework formally tests patient-specific pathway rewiring: edge- and pathway-level tests with FDR control, label-permutation nulls, and static-model and standard-GAT controls.")
add(d, "- Rewiring magnitude is clinically anchored (Ki-67 in two independent cohorts, tumor mutational burden in LUAD) and is reported transparently where associations are not significant.")
add(d, "- Open-source, pip-installable package with reproducible pipelines covering 11 TCGA cancer types and 25 independent GEO validation cohorts; no patient-level data are redistributed.")
d.save(r"manuscript/highlights.docx")
print("highlights.docx saved")
