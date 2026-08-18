# -*- coding: utf-8 -*-
"""Append Table S4 (calibration) to manuscript/supplementary.docx (python-docx).
Reads results/calibration_results.csv; creates the table if it does not exist yet.
"""
import os, sys
import pandas as pd
from docx import Document
from docx.shared import Pt

ROOT = r"D:\TT paper\Path-AGNN-Cox Pathway-Constrained Adaptive Graph Neural Network for Interpretable Survival Analysis"
CSV = os.path.join(ROOT, "results", "calibration_results.csv")
DOCX = os.path.join(ROOT, "manuscript", "supplementary.docx")

df = pd.read_csv(CSV)
# canonical column names
cols = ["dataset", "setting", "cohort", "model", "n", "events", "slope", "slope_ci_low", "slope_ci_high", "cal_mae"]
df = df[cols].copy()
df["slope"] = df["slope"].map(lambda v: "%.2f" % v if pd.notna(v) else "\u2014")
df["slope_ci_low"] = df["slope_ci_low"].map(lambda v: "%.2f" % v if pd.notna(v) else "\u2014")
df["slope_ci_high"] = df["slope_ci_high"].map(lambda v: "%.2f" % v if pd.notna(v) else "\u2014")
df["cal_mae"] = df["cal_mae"].map(lambda v: "%.3f" % v if pd.notna(v) else "\u2014")
df["n"] = df["n"].astype(int); df["events"] = df["events"].astype(int)
df["setting"] = df["setting"].map({"internal": "Internal CV", "external": "External transfer"})
df.columns = ["Dataset", "Setting", "Cohort", "Model", "N", "Events", "Slope", "95% CI low", "95% CI high", "Calibration MAE"]

doc = Document(DOCX)
# check whether Table S4 already exists (search text)
exists = any("Calibration of the risk score" in p.text for p in doc.paragraphs)
if exists:
    print("Table S4 already present; skipping")
    sys.exit(0)

doc.add_heading("Table S4. Calibration of the risk score.", level=2)
p = doc.add_paragraph()
r = p.add_run("Calibration slope from a univariate Cox regression of the standardized risk score; the ideal value is 1. "
              "Calibration MAE is the mean absolute deviation between model-predicted and Kaplan\u2013Meier survival across "
              "risk tertiles at the 25th, 50th and 75th percentiles of follow-up time. Internal rows use out-of-fold risk scores "
              "from stratified 5-fold cross-validation; external rows transfer models trained on the full TCGA cohort without fine-tuning.")
r.italic = True
r.font.size = Pt(9)

table = doc.add_table(rows=1, cols=len(df.columns))
table.style = "Table Grid"
hdr = table.rows[0].cells
for j, c in enumerate(df.columns):
    hdr[j].text = c
for _, row in df.iterrows():
    cells = table.add_row().cells
    for j, v in enumerate(row.values):
        cells[j].text = str(v)

doc.save(DOCX)
print("Table S4 appended with", len(df), "rows")