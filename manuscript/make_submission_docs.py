# -*- coding: utf-8 -*-
"""Generate submission Word documents: title page, highlights, data availability, cover letter."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(r"D:\TT paper\Path-AGNN-Cox Pathway-Constrained Adaptive Graph Neural Network for Interpretable Survival Analysis")
OUT = ROOT / "manuscript"

TITLE = ("Path-AGNN-Cox: a reproducible statistical framework for testing "
         "patient-specific pathway rewiring in cancer survival analysis")
AUTHORS = ("Zhipeng Wang[1,*], Luning Wang[2,*], Changsong Wang[1,†], Pengli Zhai[3], "
           "Zejun Liu[1], Hui Feng[1], Hongmei Liu[1], Qian Hou[1], Ming Guo[1]")
AFFIL = ("1 Department of TCM, Zhongda Hospital, Southeast University, China; "
         "2 Department of Rehabilitation Medicine, Zhongda Hospital, Southeast University, China; "
         "3 Jiangbei Campus, Jiangsu Provincial Traditional Chinese Medicine Hospital, China")
CORR = ("Corresponding author: Changsong Wang, Department of TCM, Zhongda Hospital, Southeast University, "
        "China. Email: 101005664@seu.edu.cn")

def new_doc():
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Times New Roman"
    st.font.size = Pt(11)
    return doc

def para(doc, text, bold=False, align=None, size=None, space_after=6):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    if size:
        r.font.size = Pt(size)
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p

# 1) title page
doc = new_doc()
para(doc, "Title Page", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
para(doc, TITLE, bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
para(doc, "Authors: " + AUTHORS, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
para(doc, "Affiliations: " + AFFIL, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
para(doc, CORR, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
para(doc, "Keywords", bold=True, space_after=4)
para(doc, "survival analysis; graph neural network; pathway rewiring; KEGG; interpretability; "
          "patient-specific graphs; reproducibility", space_after=12)
para(doc, "Manuscript statistics", bold=True, space_after=4)
for line in ["- Abstract: ~300 words", "- Main text: ~5,000 words (Introduction to Conclusion)",
             "- Figures: 8 (Figures 1-8)", "- Tables: 9 (Tables 1-9)",
             "- Supplementary: none (all methods and tables integrated in the main text)"]:
    para(doc, line, space_after=2)
para(doc, "Conflict of interest statement", bold=True, space_after=4)
para(doc, "The authors declare no competing financial interests.", space_after=12)
para(doc, "Code and data availability", bold=True, space_after=4)
para(doc, "All code is publicly available at https://github.com/ZhongdaHospital-SEU/Path-AGNN-Cox (MIT license) and installable via PyPI as `path-agnn-cox` (pip install path-agnn-cox); a versioned archive is available at Zenodo (https://doi.org/10.5281/zenodo.22030045). Processed datasets and full reproducible pipelines are provided; raw TCGA and GEO data are publicly available from their original sources.", space_after=12)
doc.save(OUT / "title_page.docx")
print("title_page.docx OK")

# 2) highlights
doc = new_doc()
para(doc, "Highlights", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
for h in [
    "Path-AGNN-Cox learns per-patient pathway graphs within KEGG pathway modules and links them to survival through a Cox objective.",
    "A reproducible statistical framework formally tests patient-specific pathway rewiring: edge/pathway-level tests with FDR control, label-permutation nulls, static and standard-GAT negative controls.",
    "Rewiring magnitude is clinically anchored (Ki-67 in two independent cohorts, TMB in LUAD) and is reported transparently where associations are not significant (external GEO replication, IMvigor210 response).",
    "Open-source, pip-installable package with reproducible pipelines covering 11 TCGA cancer types and 25 independent GEO validation cohorts.",
]:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(h)
doc.save(OUT / "highlights.docx")
print("highlights.docx OK")

# 3) data availability
doc = new_doc()
para(doc, "Data Availability Statement", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
for text in [
    "All data used in this study are publicly available from their original sources. TCGA level-3 "
    "RNA-seq and clinical data were obtained from the GDC data portal (https://portal.gdc.cancer.gov). "
    "Independent external validation cohorts were obtained from the Gene Expression Omnibus "
    "(https://www.ncbi.nlm.nih.gov/geo) under the accession numbers listed in Table 1 "
    "(25 cohorts across 11 cancer types). The IMvigor210 anti-PD-L1 cohort was obtained from the "
    "IMvigor210CoreBiologies R/Bioconductor package (https://github.com/BioInfoCloud/IMvigor210CoreBiologies).",
    "Processed training matrices, external cohort matrices, clinical annotations, and all analysis "
    "outputs required to reproduce the figures and tables of this manuscript are provided in the "
    "associated GitHub repository (https://github.com/ZhongdaHospital-SEU/Path-AGNN-Cox) and as a pip-installable Python package (`path-agnn-cox`; Zenodo archive: https://doi.org/10.5281/zenodo.22030045).",
    "No new biological data were generated in this study.",
]:
    para(doc, text, space_after=10)
doc.save(OUT / "data_availability.docx")
print("data_availability.docx OK")

# 4) cover letter (from cover_letter.md draft)
doc = new_doc()
for ln in Path(OUT / "cover_letter.md").read_text(encoding="utf-8").splitlines():
    s = ln.strip()
    if not s or s.startswith("#"):
        continue
    if s.startswith("- "):
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(s[2:])
    else:
        para(doc, s, space_after=8)
doc.save(OUT / "cover_letter.docx")
print("cover_letter.docx OK")