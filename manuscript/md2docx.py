# -*- coding: utf-8 -*-
"""Convert the rendered manuscript Markdown to DOCX (python-docx).

Usage:  python manuscript/md2docx.py
Output: manuscript/Path-AGNN-Cox_manuscript.docx

- headings -> Heading 1/2/3
- markdown tables -> Table Grid
- bullet lines -> List Bullet
- **bold** and *italic* inline formatting
- $$ latex $$ lines -> Word native equations (OMML) via latex2mathml + MML2OMML.XSL
- figure PNGs inserted (centered) above their legend lines in the
  'Figure legends' section, using results/figures/figure_manifest.json
"""
from __future__ import annotations
import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(r"D:\TT paper\Path-AGNN-Cox Pathway-Constrained Adaptive Graph Neural Network for Interpretable Survival Analysis")
MD = ROOT / "manuscript" / "Path-AGNN-Cox_manuscript.md"
OUT = ROOT / "manuscript" / "Path-AGNN-Cox_manuscript.docx"
MANIFEST = ROOT / "results" / "figures" / "figure_manifest.json"

BOLD = re.compile(r"\*\*(.+?)\*\*")
ITAL = re.compile(r"(?<!\*)\*([^*\[\]\n]+?)\*(?!\*)")

CITATION = re.compile(r'\[(\d+(?:\s*[,\-]\s*\d+)*)\]')
REFERENCE_LINE = re.compile(r'^(\d{1,2})\.\s+(.*)$')


def _add_reference_field(par, reference_number):
    '''Insert a Word REF field that jumps to the numbered-reference bookmark.'''
    begin_run = par.add_run()
    begin = OxmlElement('w:fldChar')
    begin.set(qn('w:fldCharType'), 'begin')
    begin_run._r.append(begin)
    instr_run = par.add_run()
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = f' REF ref{reference_number} \\h '
    instr_run._r.append(instr)
    separate_run = par.add_run()
    separate = OxmlElement('w:fldChar')
    separate.set(qn('w:fldCharType'), 'separate')
    separate_run._r.append(separate)
    par.add_run(str(reference_number))
    end_run = par.add_run()
    end = OxmlElement('w:fldChar')
    end.set(qn('w:fldCharType'), 'end')
    end_run._r.append(end)


def _add_citation_group(par, bracketed_numbers):
    numbers = []
    for part in bracketed_numbers.split(','):
        boundaries = [int(value.strip()) for value in part.split('-')]
        if len(boundaries) == 1:
            numbers.extend(boundaries)
        elif len(boundaries) == 2 and boundaries[0] <= boundaries[1]:
            numbers.extend(range(boundaries[0], boundaries[1] + 1))
        else:
            return False
    if not numbers or any(number < 1 or number > 46 for number in numbers):
        return False
    par.add_run('[')
    for index, number in enumerate(numbers):
        if index:
            par.add_run(',')
        _add_reference_field(par, number)
    par.add_run(']')
    return True


CODE = re.compile(r"`([^`\n]+?)`")

def add_runs(par, text, base_bold=False):
    """Add text, using monospace for inline code spans, then **bold** and *italic*."""
    pos = 0
    for m in CODE.finditer(text):
        if m.start() > pos:
            add_runs_md(par, text[pos:m.start()], base_bold)
        r = par.add_run(m.group(1))
        r.font.name = "Consolas"
        r.font.size = Pt(9.5)
        r.bold = base_bold
        pos = m.end()
    if pos < len(text):
        add_runs_md(par, text[pos:], base_bold)

def add_runs_md(par, text, base_bold=False):
    """Add text to paragraph, honoring **bold** and *italic* markers."""
    pos = 0
    for m in BOLD.finditer(text):
        if m.start() > pos:
            add_runs_plain(par, text[pos:m.start()], base_bold)
        r = par.add_run(m.group(1)); r.bold = True
        pos = m.end()
    if pos < len(text):
        add_runs_plain(par, text[pos:], base_bold)

UNDER = re.compile(r"(?<!\w)_([^_\n]+?)_(?!\w)")

def add_runs_plain(par, text, base_bold=False):
    pos = 0
    for citation in CITATION.finditer(text):
        if citation.start() > pos:
            add_runs_underscore(par, text[pos:citation.start()], base_bold)
        if not _add_citation_group(par, citation.group(1)):
            add_runs_underscore(par, citation.group(0), base_bold)
        pos = citation.end()
    if pos < len(text):
        add_runs_underscore(par, text[pos:], base_bold)


def add_runs_underscore(par, text, base_bold=False):
    pos = 0
    for m in UNDER.finditer(text):
        if m.start() > pos:
            _add_ital(par, text[pos:m.start()], base_bold)
        r = par.add_run(m.group(1)); r.italic = True; r.bold = base_bold
        pos = m.end()
    if pos < len(text):
        _add_ital(par, text[pos:], base_bold)

def _add_ital(par, text, base_bold=False):
    pos = 0
    for m in ITAL.finditer(text):
        if m.start() > pos:
            r = par.add_run(text[pos:m.start()]); r.bold = base_bold
        r = par.add_run(m.group(1)); r.italic = True; r.bold = base_bold
        pos = m.end()
    if pos < len(text):
        r = par.add_run(text[pos:]); r.bold = base_bold

def parse_table(block):
    rows = []
    for ln in block.strip().splitlines():
        ln = ln.strip()
        if not ln.startswith("|"):
            continue
        cells = [c.strip() for c in ln.strip("|").split("|")]
        if set("".join(cells)) <= set("-: "):
            continue
        rows.append(cells)
    return rows

from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import latex2mathml.converter
from lxml import etree

_MML2OMML = Path(r"C:\Program Files (x86)\Microsoft Office\Office12\MML2OMML.XSL")
_equation_xslt = etree.XSLT(etree.parse(str(_MML2OMML))) if _MML2OMML.exists() else None


def _latex_to_omml(latex):
    """Convert a LaTeX fragment to a Word native equation (m:oMath) element."""
    mml = latex2mathml.converter.convert(latex)
    root = etree.fromstring(mml.encode("utf-8"))
    return _equation_xslt(root).getroot()


def _border(el, tag, sz):
    e = OxmlElement(tag)
    e.set(qn("w:val"), "single")
    e.set(qn("w:sz"), str(sz))
    e.set(qn("w:space"), "0")
    e.set(qn("w:color"), "000000")
    el.append(e)


def style_booktabs(tbl):
    """Three-line table: top rule, header bottom rule, bottom rule; no vertical lines."""
    nrow = len(tbl.rows)
    for ri, row in enumerate(tbl.rows):
        for cell in row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            tcB = tcPr.find(qn("w:tcBorders"))
            if tcB is None:
                tcB = OxmlElement("w:tcBorders")
                tcPr.append(tcB)
            for tag in ("w:top", "w:left", "w:bottom", "w:right", "w:insideH", "w:insideV"):
                el = tcB.find(qn(tag))
                if el is not None:
                    tcB.remove(el)
            if ri == 0:
                _border(tcB, "w:top", 12)
                _border(tcB, "w:bottom", 6)
            if ri == nrow - 1:
                _border(tcB, "w:bottom", 12)


def main():
    text = MD.read_text(encoding="utf-8")
    lines = text.splitlines()
    man = json.loads(MANIFEST.read_text(encoding="utf-8")) if MANIFEST.exists() else {}
    png_by_num = {}
    for fname, meta in man.items():
        m = re.match(r"Figure(\d+)", fname)
        if m:
            png = ROOT / "results" / "figures" / (Path(meta["file"]).stem + ".png")
            if png.exists():
                png_by_num[int(m.group(1))] = png
    doc = Document()
    in_references = False
    bookmark_id = 1000
    # base style
    st = doc.styles["Normal"]
    st.font.name = "Times New Roman"
    st.font.size = Pt(10.5)
    i = 0
    while i < len(lines):
        ln = lines[i]
        s = ln.strip()
        if not s:
            i += 1
            continue
        if s == "---":
            i += 1
            continue
        # code fences: emit verbatim, no markdown formatting
        if s.startswith("```"):
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                par = doc.add_paragraph()
                r = par.add_run(lines[i])
                r.font.name = "Consolas"
                r.font.size = Pt(9.5)
                par.paragraph_format.space_after = Pt(0)
                i += 1
            i += 1  # skip closing fence
            continue
        # tables: detect a block of consecutive lines starting with |
        if s.startswith("|") and i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
            j = i
            while j < len(lines) and lines[j].strip().startswith("|"):
                j += 1
            rows = parse_table("\n".join(lines[i:j]))
            if rows:
                ncol = max(len(r) for r in rows)
                tbl = doc.add_table(rows=len(rows), cols=ncol)
                tbl.style = "Normal Table"
                style_booktabs(tbl)
                for ri, r in enumerate(rows):
                    for ci in range(ncol):
                        cell = tbl.cell(ri, ci)
                        cell.text = ""
                        par = cell.paragraphs[0]
                        add_runs(par, r[ci] if ci < len(r) else "")
                i = j
            else:
                i += 1
            continue
        # headings: plain bold only, no Heading styles or other formatting
        m = re.match(r"^(#{1,3})\s+(.*)$", s)
        if m:
            lvl = len(m.group(1))
            if m.group(2) == 'References':
                in_references = True
            par = doc.add_paragraph()
            add_runs(par, m.group(2))
            for r in par.runs:
                r.bold = True
            par.paragraph_format.space_before = Pt(6)
            par.paragraph_format.space_after = Pt(3)
            i += 1
            continue
        # native equations: $$ latex $$ -> Word OMML equation
        m_eq = re.match(r"^\$\$(.+)\$\$\s*$", s)
        if m_eq:
            par = doc.add_paragraph()
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                par._p.append(_latex_to_omml(m_eq.group(1).strip()))
            except Exception:
                r = par.add_run(m_eq.group(1).strip())
                r.italic = True
            i += 1
            continue
        # inline figure image line -> centered PNG (SVG is the editable source)
        mi = re.match(r"^!\[Figure\s*(\d+)\]\(([^)]+)\)$", s)
        if mi:
            num = int(mi.group(1))
            if num in png_by_num:
                par = doc.add_paragraph()
                par.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = par.add_run()
                run.add_picture(str(png_by_num[num]), width=Inches(6.0))
            i += 1
            continue
        # figure legend line -> paragraph with bold lead
        mf = re.match(r"^- \*\*Figure\s+(\d+)([A-Z]?)\.\s+(.*)$", s)
        if mf:
            num = int(mf.group(1))
            par = doc.add_paragraph()
            r = par.add_run("Figure %d%s. " % (num, mf.group(2))); r.bold = True
            rest = mf.group(3)
            idx = rest.find('**')
            if idx != -1:
                lead = rest[:idx].rstrip()
                tail = rest[idx + 2:].strip()
                if lead:
                    r2 = par.add_run(lead); r2.bold = True
                add_runs(par, (' ' + tail) if tail else '')
            else:
                add_runs(par, rest)
            par.paragraph_format.space_after = Pt(10)
            i += 1
            continue
        # bullet
        if s.startswith("- "):
            par = doc.add_paragraph(style="List Bullet")
            add_runs(par, s[2:])
            i += 1
            continue
        # reference paragraphs receive anchors used by clickable REF fields.
        if s == '## References':
            in_references = True
        reference_match = REFERENCE_LINE.match(s) if in_references else None
        if reference_match:
            par = doc.add_paragraph()
            number_run = par.add_run(reference_match.group(1))
            start = OxmlElement('w:bookmarkStart')
            start.set(qn('w:id'), str(bookmark_id))
            start.set(qn('w:name'), 'ref' + str(int(reference_match.group(1))))
            end = OxmlElement('w:bookmarkEnd')
            end.set(qn('w:id'), str(bookmark_id))
            number_run._r.addprevious(start)
            number_run._r.addnext(end)
            bookmark_id += 1
            par.add_run('. ')
            add_runs(par, reference_match.group(2))
            i += 1
            continue
        # ordinary paragraph: join soft-wrapped source lines (markdown single newline)
        joined = [s]
        j = i + 1
        while j < len(lines):
            nxt = lines[j].strip()
            if not nxt:
                break
            if re.match(r"^(#{1,3}\s+|---+$|\||\$\$|!\[)", nxt):
                break
            if nxt.startswith("- "):
                break
            if nxt.startswith('```'):
                break
            if in_references and REFERENCE_LINE.match(nxt):
                break
            if REFERENCE_LINE.match(nxt):
                break
            joined.append(nxt)
            j += 1
        par = doc.add_paragraph()
        add_runs(par, ' '.join(joined))
        i = j
        continue
    doc.save(OUT)
    print("saved:", OUT)
    print("figures embedded:", len(png_by_num))

if __name__ == "__main__":
    main()
